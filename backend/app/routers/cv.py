import logging
import os
import io
import uuid
import time
from datetime import datetime
from typing import Optional, List

from fastapi import APIRouter, Depends, HTTPException, Query, status, Request, BackgroundTasks, File, UploadFile
from sqlalchemy import select, func, or_
from sqlalchemy.ext.asyncio import AsyncSession
from PIL import Image

from app.routers.deps import get_db, get_current_user, get_house_context, has_permission
from app.schemas.cv import CVCreate, CVUpdate, CVSchema
from app.models.cv import CV, generate_slug
from app.models.user import User
from app.utils.activity_logger import log_activity
from app.utils.access_control import is_admin_user
from app.utils.validation import validate_image, MAX_FILE_SIZE

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/cv", tags=["cv"])


def get_user_house_ids(current_user: User) -> List[int]:
    if is_admin_user(current_user):
        return []
    return [h.id for h in current_user.houses]


async def resolve_cv(db: AsyncSession, identifier: str) -> Optional[CV]:
    """Resolve CV by numeric ID or slug."""
    result = await db.execute(
        select(CV).where(
            or_(
                CV.slug == identifier,
                CV.id == (int(identifier) if identifier.isdigit() else -1),
            ),
            CV.is_deleted == False,
        )
    )
    return result.scalar_one_or_none()


@router.get("")
async def list_cvs(
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.view")),
    house_context: Optional[int] = Depends(get_house_context),
    page: int = Query(1, ge=1),
    per_page: int = Query(20, ge=1, le=100),
    search: Optional[str] = Query(None),
    sort_by: str = Query("created_at"),
    sort_order: str = Query("desc"),
):
    is_admin = is_admin_user(current_user)
    user_house_ids = get_user_house_ids(current_user)

    query = select(CV).where(CV.is_deleted == False)

    if not is_admin and user_house_ids:
        query = query.where(CV.house_id.in_(user_house_ids))

    if house_context:
        if is_admin or house_context in user_house_ids:
            query = query.where(CV.house_id == house_context)

    if search:
        search_filter = or_(
            CV.name.ilike(f"%{search}%"),
            CV.mobile.ilike(f"%{search}%"),
            CV.nid_number.ilike(f"%{search}%"),
        )
        query = query.where(search_filter)

    total_count = await db.scalar(select(func.count()).select_from(query.subquery()))

    sort_column = getattr(CV, sort_by, CV.created_at)
    if sort_order == "asc":
        query = query.order_by(sort_column.asc())
    else:
        query = query.order_by(sort_column.desc())

    query = query.offset((page - 1) * per_page).limit(per_page)
    result = await db.execute(query)
    cvs = result.scalars().all()

    total_pages = max(1, (total_count + per_page - 1) // per_page) if total_count else 1

    return {
        "success": True,
        "data": [CVSchema.model_validate(cv).model_dump() for cv in cvs],
        "pagination": {
            "page": page,
            "per_page": per_page,
            "total": total_count or 0,
            "total_pages": total_pages,
            "has_next": page < total_pages,
            "has_prev": page > 1,
        },
    }


@router.get("/{identifier}")
async def get_cv(
    identifier: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.view")),
):
    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")
    return {"success": True, "data": CVSchema.model_validate(cv).model_dump()}


@router.post("", status_code=status.HTTP_201_CREATED)
async def create_cv(
    data: CVCreate,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.create")),
    house_context: Optional[int] = Depends(get_house_context),
):
    cv = CV(
        user_id=current_user.id,
        house_id=house_context or (current_user.houses[0].id if current_user.houses else None),
        **data.model_dump(exclude_unset=True),
    )
    db.add(cv)
    await db.flush()
    cv.slug = generate_slug(cv.name, cv.id)

    await db.commit()
    await db.refresh(cv)

    await log_activity(
        db=db,
        user_id=current_user.id,
        user_name=current_user.name,
        module="cv",
        action="create",
        record_id=cv.id,
        record_identifier=cv.name,
        new_values=data.model_dump(),
        request=request,
    )

    return {"success": True, "data": CVSchema.model_validate(cv).model_dump()}


@router.put("/{identifier}")
async def update_cv(
    identifier: str,
    data: CVUpdate,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.edit")),
    house_context: Optional[int] = Depends(get_house_context),
):
    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    user_house_ids = get_user_house_ids(current_user)
    if not is_admin_user(current_user) and cv.house_id and cv.house_id not in user_house_ids:
        raise HTTPException(status_code=403, detail="Access denied")

    old_values = CVSchema.model_validate(cv).model_dump()
    update_data = data.model_dump(exclude_unset=True)
    for field, value in update_data.items():
        setattr(cv, field, value)
    cv.updated_at = datetime.utcnow()

    if "name" in update_data and update_data["name"] != old_values.get("name"):
        cv.slug = generate_slug(cv.name, cv.id)

    await db.commit()
    await db.refresh(cv)

    await log_activity(
        db=db,
        user_id=current_user.id,
        user_name=current_user.name,
        module="cv",
        action="edit",
        record_id=cv.id,
        record_identifier=cv.name,
        old_values=old_values,
        new_values=update_data,
        request=request,
    )

    return {"success": True, "data": CVSchema.model_validate(cv).model_dump()}


@router.delete("/{identifier}")
async def delete_cv(
    identifier: str,
    request: Request,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.delete")),
):
    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    old_values = CVSchema.model_validate(cv).model_dump()
    cv.is_deleted = True
    cv.deleted_at = datetime.utcnow()
    cv.deleted_by = current_user.id

    await db.commit()

    await log_activity(
        db=db,
        user_id=current_user.id,
        user_name=current_user.name,
        module="cv",
        action="delete",
        record_id=cv.id,
        record_identifier=cv.name,
        old_values=old_values,
        request=request,
    )

    return {"success": True, "message": "CV deleted successfully"}


@router.post("/{identifier}/upload-photo")
async def upload_cv_photo(
    identifier: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.edit")),
):
    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    filename = file.filename or "photo.jpg"
    if not validate_image(filename):
        raise HTTPException(status_code=400, detail="Invalid file type. Only JPG, JPEG, and PNG files are allowed.")

    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50 MB.")

    image = Image.open(io.BytesIO(contents))
    if image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
    max_size = (400, 400)
    image.thumbnail(max_size, Image.Resampling.LANCZOS)
    buffer = io.BytesIO()
    image.save(buffer, format="JPEG", quality=85, optimize=True)

    file_name = f"cv_photo_{uuid.uuid4()}.jpg"
    file_path = f"uploads/cv/{file_name}"
    os.makedirs("uploads/cv", exist_ok=True)
    with open(file_path, "wb") as f:
        f.write(buffer.getvalue())

    if cv.photo_url:
        old_path = cv.photo_url.lstrip("/")
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass

    cv.photo_url = f"/uploads/cv/{file_name}"
    await db.commit()
    await db.refresh(cv)
    return {"url": cv.photo_url}


@router.post("/{identifier}/upload-signature")
async def upload_cv_signature(
    identifier: str,
    file: UploadFile = File(...),
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.edit")),
):
    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    filename = file.filename or "signature.png"
    if not validate_image(filename):
        raise HTTPException(status_code=400, detail="Invalid file type. Only JPG, JPEG, and PNG files are allowed.")

    contents = await file.read()
    if len(contents) > MAX_FILE_SIZE:
        raise HTTPException(status_code=400, detail="File too large. Maximum size is 50 MB.")

    image = Image.open(io.BytesIO(contents))
    if image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
    max_size = (300, 150)
    image.thumbnail(max_size, Image.Resampling.LANCZOS)
    buffer = io.BytesIO()
    image.save(buffer, format="PNG", optimize=True)

    file_name = f"cv_sig_{uuid.uuid4()}.png"
    file_path = f"uploads/cv/{file_name}"
    os.makedirs("uploads/cv", exist_ok=True)
    with open(file_path, "wb") as f:
        f.write(buffer.getvalue())

    if cv.signature_url:
        old_path = cv.signature_url.lstrip("/")
        if os.path.exists(old_path):
            try:
                os.remove(old_path)
            except OSError:
                pass

    cv.signature_url = f"/uploads/cv/{file_name}"
    await db.commit()
    await db.refresh(cv)
    return {"url": cv.signature_url}


def cleanup_temp(path: str):
    import os
    try:
        os.unlink(path)
    except Exception:
        pass


@router.get("/{identifier}/export/word")
async def export_cv_word(
    identifier: str,
    background_tasks: BackgroundTasks,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.view")),
):
    from fastapi.responses import FileResponse
    import tempfile

    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="python-docx is not installed on the server. Run: pip install python-docx and restart the backend.",
        )

    doc = Document()

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Aptos Narrow'
    font.size = Pt(12)

    p = doc.add_paragraph()
    run = p.add_run("CURRICULUM VITAE")
    run.bold = True
    run.font.size = Pt(22)
    run.font.name = 'Aptos Narrow'

    photo_path = None
    if cv.photo_url:
        candidate = cv.photo_url.lstrip("/")
        if os.path.exists(candidate):
            photo_path = candidate

    if photo_path:
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        cell_left = table.cell(0, 0)
        cell_left.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = cell_left.paragraphs[0].add_run(cv.name.upper() if cv.name else "")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Aptos Narrow'
        if cv.care_of:
            p = cell_left.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"C/O: {cv.care_of}")
            run.font.name = 'Aptos Narrow'
        if cv.mobile:
            p = cell_left.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(f"Mobile: {cv.mobile}")
            run.font.name = 'Aptos Narrow'

        cell_right = table.cell(0, 1)
        cell_right.width = Inches(1.5)
        run = cell_right.paragraphs[0].add_run()
        run.add_picture(photo_path, width=Inches(1.2), height=Inches(1.5))
    else:
        p = doc.add_paragraph()
        run = p.add_run(cv.name.upper() if cv.name else "")
        run.bold = True
        run.font.size = Pt(14)
        run.font.name = 'Aptos Narrow'

        if cv.care_of:
            p = doc.add_paragraph()
            run = p.add_run(f"C/O: {cv.care_of}")
            run.font.name = 'Aptos Narrow'
        if cv.mobile:
            p = doc.add_paragraph()
            run = p.add_run(f"Mobile: {cv.mobile}")
            run.font.name = 'Aptos Narrow'

    def add_section_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(12)
        run.font.name = 'Aptos Narrow'

    def add_info_row(label, value):
        if value:
            p = doc.add_paragraph()
            run = p.add_run(f"{label}: ")
            run.bold = True
            run.font.name = 'Aptos Narrow'
            run = p.add_run(str(value))
            run.font.name = 'Aptos Narrow'

    doc.add_paragraph()
    add_section_header("PERSONAL INFORMATION")
    add_info_row("Name", cv.name)
    add_info_row("Father's Name", cv.fathers_name)
    add_info_row("Mother's Name", cv.mothers_name)
    add_info_row("Permanent Address", cv.permanent_address)
    if cv.date_of_birth:
        add_info_row("Date of Birth", cv.date_of_birth.strftime("%d/%m/%Y"))
    add_info_row("NID Number", cv.nid_number)
    add_info_row("Nationality", cv.nationality)
    add_info_row("Religion", cv.religion)
    add_info_row("Marital Status", cv.marital_status)
    add_info_row("Blood Group", cv.blood_group)

    if cv.educational_qualifications:
        doc.add_paragraph()
        add_section_header("EDUCATIONAL QUALIFICATION")
        for edu in cv.educational_qualifications:
            p = doc.add_paragraph()
            run = p.add_run(f"{edu.get('degree', '')} ({edu.get('group_subject', '')})")
            run.bold = True
            run.font.name = 'Aptos Narrow'
            p = doc.add_paragraph()
            run = p.add_run(f"Board: {edu.get('board', '')} | Result: {edu.get('result', '')}")
            run.font.name = 'Aptos Narrow'
            p = doc.add_paragraph()
            run = p.add_run(f"Institution: {edu.get('institution', '')}")
            run.font.name = 'Aptos Narrow'
            p = doc.add_paragraph()
            run = p.add_run(f"Passing Year: {edu.get('passing_year', '')}")
            run.font.name = 'Aptos Narrow'

    if cv.professional_experiences:
        doc.add_paragraph()
        add_section_header("PROFESSIONAL EXPERIENCE")
        for exp in cv.professional_experiences:
            p = doc.add_paragraph()
            run = p.add_run(f"Institution: {exp.get('institution', '')}")
            run.font.name = 'Aptos Narrow'
            add_info_row("Designation", exp.get('designation', ''))
            add_info_row("Duration", exp.get('duration', ''))
            add_info_row("Responsibilities", "")
            for resp in exp.get('responsibilities', []):
                p = doc.add_paragraph()
                run = p.add_run(resp)
                run.font.name = 'Aptos Narrow'

    if cv.language_proficiency:
        doc.add_paragraph()
        add_section_header("LANGUAGE PROFICIENCY")
        p = doc.add_paragraph()
        run = p.add_run(cv.language_proficiency)
        run.font.name = 'Aptos Narrow'

    doc.add_paragraph()
    add_section_header("DECLARATION")
    decl_text = cv.declaration_text or "I, the undersigned, certify that all information stated herein is true and correct."
    p = doc.add_paragraph()
    run = p.add_run(decl_text)
    run.font.name = 'Aptos Narrow'

    doc.add_paragraph()
    sig_path = None
    if cv.signature_url:
        candidate = cv.signature_url.lstrip("/")
        if os.path.exists(candidate):
            sig_path = candidate
    if sig_path:
        s = doc.add_paragraph()
        s.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = s.add_run()
        run.add_picture(sig_path, width=Inches(2), height=Inches(0.5))
    signature_text = cv.signature_name or cv.name or ""
    p = doc.add_paragraph()
    run = p.add_run("Signature: ")
    run.bold = True
    run.font.name = 'Aptos Narrow'
    run = p.add_run("___________________")
    run.font.name = 'Aptos Narrow'
    if signature_text:
        p = doc.add_paragraph()
        run = p.add_run(signature_text)
        run.bold = True
        run.font.name = 'Aptos Narrow'
    if cv.declaration_date:
        p = doc.add_paragraph()
        run = p.add_run("Date: ")
        run.bold = True
        run.font.name = 'Aptos Narrow'
        run = p.add_run(f"___________ ({cv.declaration_date.strftime('%d/%m/%Y')})")
        run.font.name = 'Aptos Narrow'

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    background_tasks.add_task(cleanup_temp, tmp.name)

    filename = f"CV_{cv.slug or cv.name.replace(' ', '_')}.docx"

    return FileResponse(
        path=tmp.name,
        filename=filename,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{identifier}/export/pdf")
async def export_cv_pdf(
    identifier: str,
    db: AsyncSession = Depends(get_db),
    current_user: User = Depends(has_permission("cv.view")),
):
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import mm
        from reportlab.lib import colors
    except ImportError:
        raise HTTPException(
            status_code=500,
            detail="reportlab is not installed. Run: pip install reportlab",
        )

    from fastapi.responses import FileResponse
    import tempfile

    cv = await resolve_cv(db, identifier)
    if not cv:
        raise HTTPException(status_code=404, detail="CV not found")

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
    doc = SimpleDocTemplate(tmp.name, pagesize=A4,
                            topMargin=20*mm, bottomMargin=20*mm,
                            leftMargin=25*mm, rightMargin=25*mm)
    styles = getSampleStyleSheet()
    story = []

    title_style = ParagraphStyle('CVTitle', parent=styles['Title'],
                                  fontSize=20, leading=24, spaceAfter=12,
                                  alignment=1, fontName='Times-Bold')
    name_style = ParagraphStyle('CVName', parent=styles['Normal'],
                                 fontSize=14, leading=16, alignment=2,
                                 fontName='Times-Bold', spaceAfter=4)
    normal = ParagraphStyle('CVNormal', parent=styles['Normal'],
                            fontSize=11, leading=14, fontName='Times-Roman')
    section_style = ParagraphStyle('CVSection', parent=styles['Normal'],
                                    fontSize=13, leading=15, spaceBefore=10,
                                    spaceAfter=4, fontName='Times-Bold',
                                    underlineWidth=0.5)
    bold = ParagraphStyle('CVBold', parent=styles['Normal'],
                           fontSize=11, leading=14, fontName='Times-Bold')

    story.append(Paragraph("CURRICULUM VITAE", title_style))
    story.append(Paragraph(cv.name.upper(), name_style))
    if cv.care_of:
        story.append(Paragraph(f"C/O: {cv.care_of}", name_style))
    if cv.mobile:
        story.append(Paragraph(f"Mobile: {cv.mobile}", name_style))

    def add_section(text):
        story.append(Paragraph(text, section_style))

    def add_row(label, value):
        if value:
            story.append(Paragraph(f"<b>{label}:</b> {value}", normal))

    story.append(Spacer(1, 6))
    add_section("PERSONAL INFORMATION")
    add_row("Name", cv.name)
    add_row("Father's Name", cv.fathers_name)
    add_row("Mother's Name", cv.mothers_name)
    add_row("Permanent Address", cv.permanent_address)
    if cv.date_of_birth:
        add_row("Date of Birth", cv.date_of_birth.strftime("%d/%m/%Y"))
    add_row("NID Number", cv.nid_number)
    add_row("Nationality", cv.nationality)
    add_row("Religion", cv.religion)
    add_row("Marital Status", cv.marital_status)
    add_row("Blood Group", cv.blood_group)

    if cv.educational_qualifications:
        story.append(Spacer(1, 6))
        add_section("EDUCATIONAL QUALIFICATION")
        edu_data = [["Degree", "Subject", "Board", "Result", "Institution", "Year"]]
        for edu in cv.educational_qualifications:
            edu_data.append([
                edu.get('degree', ''), edu.get('group_subject', ''),
                edu.get('board', ''), edu.get('result', ''),
                edu.get('institution', ''), str(edu.get('passing_year', '')),
            ])
        t = Table(edu_data, colWidths=[60, 55, 50, 45, 100, 40])
        t.setStyle(TableStyle([
            ('GRID', (0, 0), (-1, -1), 0.5, colors.black),
            ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.9)),
            ('FONTNAME', (0, 0), (-1, 0), 'Times-Bold'),
            ('FONTNAME', (0, 1), (-1, -1), 'Times-Roman'),
            ('FONTSIZE', (0, 0), (-1, -1), 9),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
            ('TOPPADDING', (0, 0), (-1, -1), 3),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 3),
        ]))
        story.append(t)

    if cv.professional_experiences:
        story.append(Spacer(1, 6))
        add_section("PROFESSIONAL EXPERIENCE")
        for exp in cv.professional_experiences:
            story.append(Paragraph(f"<b>Institution:</b> {exp.get('institution', '')}", normal))
            story.append(Paragraph(f"<b>Designation:</b> {exp.get('designation', '')}", normal))
            story.append(Paragraph(f"<b>Duration:</b> {exp.get('duration', '')}", normal))
            for resp in exp.get('responsibilities', []):
                story.append(Paragraph(f"• {resp}", normal))

    if cv.language_proficiency:
        story.append(Spacer(1, 6))
        add_section("LANGUAGE PROFICIENCY")
        story.append(Paragraph(cv.language_proficiency, normal))

    story.append(Spacer(1, 6))
    add_section("DECLARATION")
    decl = cv.declaration_text or "I, the undersigned, certify that all information stated herein is true and correct."
    story.append(Paragraph(decl, normal))
    story.append(Spacer(1, 10))
    if cv.signature_name:
        story.append(Paragraph(f"Signature: ___________________{cv.signature_name}", normal))
    if cv.declaration_date:
        story.append(Paragraph(f"Date: ___________ ({cv.declaration_date.strftime('%d/%m/%Y')})", normal))

    doc.build(story)
    filename = f"CV_{cv.slug or cv.name.replace(' ', '_')}.pdf"

    return FileResponse(
        path=tmp.name,
        filename=filename,
        media_type="application/pdf",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )
